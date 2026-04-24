# -*- coding: utf-8 -*-
"""
逐章提取论文内容 v3 - 通过Heading样式识别章节结构
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
import re

doc = Document('D:/Users/MI/120599705/WPS云盘/博士毕业/【排版结果】博士论文_V14_精修版.docx')

# 找到所有Heading 1（章标题）
chapters = []
for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else "None"
    text = para.text.strip()
    if not text:
        continue
    if style_name == 'Heading 1':
        chapters.append((i, text))

print(f"找到 {len(chapters)} 个章标题（Heading 1）")
for idx, (pos, text) in enumerate(chapters):
    print(f"  [{idx}] 段落{pos}: {text[:80]}")

# 提取各章内容
for idx in range(len(chapters)):
    start_pos = chapters[idx][0]
    code = f"ch{idx}"
    title = chapters[idx][1][:50].replace(' ', '_').replace('/', '_')
    if idx + 1 < len(chapters):
        end_pos = chapters[idx + 1][0]
    else:
        end_pos = len(doc.paragraphs)
    
    # 提取该章所有段落（包括各级标题和正文）
    chapter_lines = []
    for i in range(start_pos, end_pos):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else "None"
        # 标记样式级别
        if style == 'Heading 1':
            chapter_lines.append(f"\n=== {text} ===\n")
        elif style == 'Heading 2':
            chapter_lines.append(f"\n--- {text} ---\n")
        elif style == 'Heading 3':
            chapter_lines.append(f"\n{text}\n")
        elif style == 'Heading 4':
            chapter_lines.append(f"\n{text}\n")
        elif style == 'Heading 5':
            chapter_lines.append(f"\n{text}\n")
        elif '图片标题' in style:
            chapter_lines.append(f"[图] {text}")
        elif '表格标题' in style:
            chapter_lines.append(f"[表] {text}")
        elif '参考文献' in style:
            chapter_lines.append(f"[文献] {text}")
        else:
            chapter_lines.append(text)
    
    # 保存到文件
    filename = f"_thesis_chapter{idx:02d}_{title}.txt"
    with open(filename, 'w', encoding='utf-8') as f:
        f.write('\n'.join(chapter_lines))
    
    total_chars = sum(len(line) for line in chapter_lines)
    print(f"  已保存 {filename}: {len(chapter_lines)} 行, {total_chars} 字符")

print("\n全部章节提取完成")
