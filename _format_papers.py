#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Format 5 papers for journal submission.
Applies standard Chinese academic journal formatting.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import re


def set_run_font(run, name='Times New Roman', east_asia='宋体', size=Pt(12), bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def is_heading_text(text):
    """Guess if paragraph is a heading based on content patterns."""
    t = text.strip()
    if re.match(r'^\d+([.．]\d+)*\s+', t):
        return True
    if t.startswith('表') and re.match(r'^表\d+[-－]', t):
        return True
    if t.startswith('图') and re.match(r'^图\d+[-－]', t):
        return True
    return False


def guess_heading_level(text):
    """Guess heading level from numbering pattern."""
    t = text.strip()
    # 表/图标题
    if t.startswith('表') or t.startswith('图'):
        return 4
    # 1. / 1 / 一、
    if re.match(r'^\d+([.．]\d+)*\s+', t):
        parts = re.split(r'[.．]', t.split()[0])
        return min(len(parts), 3)
    if re.match(r'^[一二三四五六七八九十]+[、．]', t):
        return 1
    return 0


def format_paper(input_path, output_path, journal_config=None):
    doc = Document(input_path)

    # Page setup: A4, margins 2.54cm top/bottom, 3.17cm left/right
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        level = guess_heading_level(text)

        # Heading formatting
        if level == 1:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            para.paragraph_format.first_line_indent = Cm(0)
            for run in para.runs:
                set_run_font(run, east_asia='黑体', size=Pt(16), bold=True)
        elif level == 2:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            para.paragraph_format.first_line_indent = Cm(0)
            for run in para.runs:
                set_run_font(run, east_asia='黑体', size=Pt(14), bold=True)
        elif level == 3:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            para.paragraph_format.first_line_indent = Cm(0)
            for run in para.runs:
                set_run_font(run, east_asia='黑体', size=Pt(12), bold=True)
        elif level == 4:
            # Table/Figure captions
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            para.paragraph_format.first_line_indent = Cm(0)
            for run in para.runs:
                set_run_font(run, east_asia='黑体', size=Pt(10.5), bold=False)
        else:
            # Body text
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.first_line_indent = Cm(0.74)  # ~2 characters
            for run in para.runs:
                set_run_font(run, east_asia='宋体', size=Pt(12), bold=run.bold)

    # Format tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        set_run_font(run, east_asia='宋体', size=Pt(10.5), bold=run.bold)

    doc.save(output_path)
    print(f"Formatted: {output_path}")


PAPERS = [
    (r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文F_fsQCA组态分析_投稿版.docx",
     r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文F_fsQCA组态分析_定稿版.docx"),
    (r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文B_数据口径幻觉_投稿版.docx",
     r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文B_数据口径幻觉_定稿版.docx"),
    (r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文D_TOPSIS-DEA绩效评估_投稿版.docx",
     r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文D_TOPSIS-DEA绩效评估_定稿版.docx"),
    (r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文A_异构平台统一采集_投稿版.docx",
     r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文A_异构平台统一采集_定稿版.docx"),
    (r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文C_4E评估框架_投稿版.docx",
     r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文C_4E评估框架_定稿版.docx"),
]

if __name__ == '__main__':
    for inp, out in PAPERS:
        try:
            format_paper(inp, out)
        except Exception as e:
            print(f"Error formatting {inp}: {e}")
