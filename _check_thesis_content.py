# -*- coding: utf-8 -*-
from docx import Document
import re

doc = Document('D:/Users/MI/120599705/WPS云盘/博士毕业/【排版结果】博士论文_V14_精修版.docx')

print('段落总数:', len(doc.paragraphs))
print('表格总数:', len(doc.tables))
print()

# 提取章节结构
print('=== 论文章节结构 ===')
heading_pattern = re.compile(r'^(第[一二三四五六七八九十]+章|第\d+章|绪论|结论|参考文献|致谢|附录|摘要|Abstract)')
section_pattern = re.compile(r'^(\d+\.\d+|\d+\.\d+\.\d+)\s')

chapter_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    style = para.style.name if para.style else ''
    
    # 检测章标题
    if heading_pattern.match(text) or ('章' in text[:15] and len(text) < 60):
        print(f'\n[章 {i}] {text[:100]}')
        chapter_count += 1
    # 检测节标题（1.1, 1.1.1等格式）
    elif section_pattern.match(text) and len(text) < 80:
        print(f'  [节 {i}] {text[:100]}')
    # 检测以数字开头的标题
    elif text[0].isdigit() and ('.' in text[:5] or ' ' in text[:10]) and len(text) < 80:
        if not text[0].isdigit() or text[1] != '.':
            continue
        print(f'  [节 {i}] {text[:100]}')

print(f'\n总共检测到约 {chapter_count} 个主要章节')

# 提取包含"开放数林"或"4E"或"fsQCA"或"DEA"的段落
print('\n\n=== 包含关键词的段落 ===')
keywords = ['开放数林', '4E', 'fsQCA', 'DEA', '数据集', '平台', '数据开放', '省级']
for kw in keywords:
    print(f'\n--- 关键词: {kw} ---')
    count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if kw in text and len(text) > 10:
            print(f'  [{i}] {text[:120]}')
            count += 1
            if count >= 5:
                print(f'  ... (还有 {sum(1 for p in doc.paragraphs if kw in p.text and len(p.text) > 10) - 5} 处)')
                break

# 提取表格说明
print('\n\n=== 表格信息 ===')
for i, table in enumerate(doc.tables[:10]):
    if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
        first_row = ' | '.join([cell.text[:20] for cell in table.rows[0].cells])
        print(f'  表格{i}: {len(table.rows)}行 x {len(table.rows[0].cells)}列 | 首行: {first_row[:80]}')

# 搜索特定章节内容
print('\n\n=== 第三章内容预览 ===')
in_ch3 = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '第三章' in text or text.startswith('3 ') or text.startswith('3.'):
        in_ch3 = True
        print(f'\n[{i}] === {text[:80]} ===')
    elif in_ch3 and text and len(text) > 20:
        print(f'  [{i}] {text[:150]}')
        # 只打印前20个内容段落
        if i > 5000:
            break
    elif in_ch3 and ('第四章' in text or text.startswith('4 ') or text.startswith('4.')):
        print(f'\n[{i}] === 进入第四章 ===')
        break
