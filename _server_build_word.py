# -*- coding: utf-8 -*-
"""
服务器端Word生成脚本
- 读取Markdown文件
- 生成精细排版的Word文档
- 图片内嵌
- 表格学术化
"""
import os
import re
import base64
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                          r'</w:tcBorders>')
    tcPr.append(tcBorders)

def add_formatted_paragraph(doc, text, style_name='Normal', alignment=None, bold=False, font_size=None, first_line_indent=None, font_name='宋体'):
    """添加格式化的段落"""
    p = doc.add_paragraph(style=style_name)
    if alignment:
        p.alignment = alignment
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if bold:
        run.font.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    
    return p

def setup_document_styles(doc):
    """设置文档样式（武大博士论文格式）"""
    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 正文样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(0.74)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

def process_markdown_to_word(doc, content, image_dir='/opt/ogd-collector-pro/static'):
    """处理Markdown内容并添加到Word文档"""
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    
    while i < len(lines):
        line = lines[i]
        
        # 代码块处理
        if line.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(80, 80, 80)
                in_code_block = False
                code_content = []
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # 空行
        if not line.strip():
            i += 1
            continue
        
        # 一级标题（章标题）
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            title = re.sub(r'第\d+章\s+第[一二三四五六七八九十]章\s*', '第', title)
            add_formatted_paragraph(doc, title, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=16, first_line_indent=None, font_name='黑体')
            i += 1
            continue
        
        # 二级标题
        if line.startswith('## ') and not line.startswith('### '):
            title = line[3:].strip()
            add_formatted_paragraph(doc, title, bold=True, font_size=14, first_line_indent=None, font_name='黑体')
            i += 1
            continue
        
        # 三级标题
        if line.startswith('### ') and not line.startswith('#### '):
            title = line[4:].strip()
            add_formatted_paragraph(doc, title, bold=True, font_size=12, first_line_indent=None, font_name='黑体')
            i += 1
            continue
        
        # 四级标题
        if line.startswith('#### '):
            title = line[5:].strip()
            add_formatted_paragraph(doc, title, bold=True, font_size=12, first_line_indent=Cm(0.74), font_name='黑体')
            i += 1
            continue
        
        # 图片引用
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            alt_text, img_path = img_match.groups()
            full_path = os.path.join(image_dir, os.path.basename(img_path))
            if not os.path.exists(full_path):
                full_path = img_path if os.path.exists(img_path) else None
            
            if full_path and os.path.exists(full_path):
                try:
                    doc.add_picture(full_path, width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    if alt_text:
                        caption_p = doc.add_paragraph()
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption_p.add_run(alt_text)
                        caption_run.font.name = 'Times New Roman'
                        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        caption_run.font.size = Pt(10.5)
                        caption_run.font.color.rgb = RGBColor(80, 80, 80)
                except Exception as e:
                    print(f'  图片插入失败: {full_path} - {e}')
            else:
                print(f'  图片不存在: {img_path}')
            i += 1
            continue
        
        # 表格处理
        if line.startswith('|') and i + 1 < len(lines) and lines[i + 1].startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            if len(table_lines) >= 2:
                rows = []
                for j, table_line in enumerate(table_lines):
                    if j == 1:
                        continue
                    cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                    if cells:
                        rows.append(cells)
                
                if rows:
                    num_cols = max(len(row) for row in rows)
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    for row_idx, row_data in enumerate(rows):
                        row = table.rows[row_idx]
                        for col_idx in range(num_cols):
                            cell = row.cells[col_idx]
                            if col_idx < len(row_data):
                                cell.text = row_data[col_idx]
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.name = 'Times New Roman'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                    run.font.size = Pt(10.5)
                            set_cell_border(cell)
                    
                    if len(table.rows) > 0:
                        for cell in table.rows[0].cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                    
                    doc.add_paragraph()
            continue
        
        # 普通段落
        text = line.strip()
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        text = re.sub(r'^>\s*', '', text)
        
        if text:
            add_formatted_paragraph(doc, text, first_line_indent=Cm(0.74))
        
        i += 1

# 主程序
print('=== 服务器端Word生成 ===')

with open('/tmp/thesis_v7.md', 'r', encoding='utf-8') as f:
    content = f.read()

print(f'读取Markdown: {len(content)} 字符')

# 创建Word文档
doc = Document()
setup_document_styles(doc)

# 封面页
add_formatted_paragraph(doc, '武汉大学', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=22, first_line_indent=None, font_name='黑体')
doc.add_paragraph()
doc.add_paragraph()
add_formatted_paragraph(doc, '博士学位论文', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=22, first_line_indent=None, font_name='黑体')
doc.add_paragraph()
doc.add_paragraph()
add_formatted_paragraph(doc, '中国省级政府数据开放平台绩效评估研究', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=18, first_line_indent=None, font_name='黑体')
add_formatted_paragraph(doc, '——基于4E框架的实证分析', alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=16, first_line_indent=None, font_name='宋体')
doc.add_paragraph()
doc.add_paragraph()
add_formatted_paragraph(doc, '研究生姓名：XXX', alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14, first_line_indent=None)
add_formatted_paragraph(doc, '指导教师：XXX 教授', alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14, first_line_indent=None)
add_formatted_paragraph(doc, '学科专业：行政管理', alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14, first_line_indent=None)
doc.add_paragraph()
doc.add_paragraph()
add_formatted_paragraph(doc, '二〇二六年四月', alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14, first_line_indent=None)

doc.add_page_break()

# 摘要
add_formatted_paragraph(doc, '摘  要', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=16, first_line_indent=None, font_name='黑体')
add_formatted_paragraph(doc, '（中文摘要内容待补充，约800-1000字）', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
doc.add_page_break()

add_formatted_paragraph(doc, 'Abstract', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=16, first_line_indent=None, font_name='黑体')
add_formatted_paragraph(doc, '(English abstract to be added, about 800-1000 words)', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
doc.add_page_break()

# 目录占位
add_formatted_paragraph(doc, '目  录', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=16, first_line_indent=None, font_name='黑体')
add_formatted_paragraph(doc, '（目录自动生成，此处为占位）', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
doc.add_page_break()

# 正文
print('处理正文...')
process_markdown_to_word(doc, content)

# 保存
output_path = '/tmp/博士论文_最终完整版_v7.docx'
doc.save(output_path)

print(f'Word生成完成: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')
