# -*- coding: utf-8 -*-
"""将Markdown转换为Word文档（v17版本）"""
import os
import sys

# 确保安装python-docx
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_MD = "docs/博士论文_最终定稿版_v10.md"
OUTPUT_DOCX = "docs/博士论文_最终完整版_v17.docx"

def md_to_docx():
    """将Markdown文件转换为Word文档"""
    doc = Document()

    # 页面设置（A4）
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.18)
    section.bottom_margin = Inches(0.98)
    section.left_margin = Inches(1.10)
    section.right_margin = Inches(0.98)

    # 默认样式
    default_style = doc.styles['Normal']
    default_style.font.name = 'Times New Roman'
    default_style.font.size = Pt(12)
    default_style.paragraph_format.line_spacing = 1.5

    # 设置中文字体
    try:
        default_style.element.rPr.rFonts.set(
            docx.oxml.ns.qn('w:eastAsia'), '宋体'
        )
    except:
        pass

    # 读取Markdown文件
    with open(INPUT_MD, 'r', encoding='utf-8') as f:
        content = f.read()

    # 解析Markdown内容并添加到文档
    lines = content.split('\n')

    for line in lines:
        line = line.rstrip('\n')
        if not line.strip():
            # 空行
            doc.add_paragraph()
            continue

        # 处理标题
        if line.startswith('###### '):
            # 6级标题
            p = doc.add_heading(line[7:], level=6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
        elif line.startswith('##### '):
            # 5级标题
            p = doc.add_heading(line[6:], level=5)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
        elif line.startswith('#### '):
            # 4级标题
            p = doc.add_heading(line[5:], level=4)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
        elif line.startswith('### '):
            # 3级标题
            p = doc.add_heading(line[3:], level=3)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(15)
        elif line.startswith('## '):
            # 2级标题
            p = doc.add_heading(line[2:], level=2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(18)
        elif line.startswith('# '):
            # 1级标题
            p = doc.add_heading(line[1:], level=1)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(22)
        elif line.startswith('**') and line.endswith('**'):
            # 粗体段落
            p = doc.add_paragraph(line.replace('**', ''))
            for run in p.runs:
                run.font.bold = True
        elif line.startswith('- ') or line.startswith('* '):
            # 无序列表
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('| '):
            # 表格行 - 跳过，因为表格需要特殊处理
            continue
        elif line.startswith('```'):
            # 代码块 - 跳过
            continue
        elif line.startswith('!['):
            # 图片 - 跳过
            continue
        else:
            # 普通段落
            p = doc.add_paragraph(line)

    # 保存文档
    doc.save(OUTPUT_DOCX)
    print(f"成功生成: {OUTPUT_DOCX}")
    return True

if __name__ == "__main__":
    md_to_docx()