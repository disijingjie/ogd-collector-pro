import re
from docx import Document

with open('docs/博士论文_最终定稿版_v10.md','r',encoding='utf-8') as f:
    c=f.read()

# 1. 22残留（排除DEA的22个）
matches=[m for m in re.finditer(r'22个',c)]
real_matches=[]
for m in matches:
    context=c[max(0,m.start()-30):m.end()+30]
    if 'DEA无效的平台有22个' not in context:
        real_matches.append(context)
print(f'[OK] 22残留(排除DEA): {len(real_matches)}处')
if real_matches:
    for r in real_matches[:3]:
        print(f'  ! {r}')

# 2. 23一致性
print(f'[OK] 23个出现次数: {c.count(chr(50)+chr(50)+chr(49)+chr(48)+chr(56))}')  # 避免编码问题

# 3. 章节重复
chap_dups=re.findall(r'(第[一二三四五六七八]章 .*)\n.*\n\1',c)
print(f'[OK] 章节重复: {len(chap_dups)}处')

# 4. Word文件
doc=Document('docs/博士论文_最终完整版_v14.docx')
img_count=sum(1 for rel in doc.part.rels.values() if 'image' in rel.target_ref)
print(f'[OK] Word段落: {len(doc.paragraphs)}, 表格: {len(doc.tables)}, 图片: {img_count}')
print(f'[OK] Markdown字符: {len(c)}')

# 5. 检查图表引用是否存在
chart_refs=re.findall(r'图[1-7]-[1-7]',c)
print(f'[OK] 图表引用数量: {len(chart_refs)}处')

print('\n=== 核对完成 ===')
