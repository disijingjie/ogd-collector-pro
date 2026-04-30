#!/usr/bin/env python3
"""将论文markdown转换为docx，支持标题、粗体、表格、引用块"""
import sys, os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='宋体', size=10.5, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def convert_markdown_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    # 设置默认中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 跳过YAML frontmatter
        if line.strip() == '---':
            if i == 0 or (i > 0 and lines[i-1].strip() == ''):
                # 找到下一个---
                j = i + 1
                while j < len(lines) and lines[j].strip() != '---':
                    j += 1
                i = j + 1
                continue

        # 标题
        if line.startswith('# '):
            p = doc.add_heading('', level=0)
            add_formatted_text(p, line[2:].strip(), '黑体', 16, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            p = doc.add_heading('', level=1)
            add_formatted_text(p, line[3:].strip(), '黑体', 14, bold=True)
        elif line.startswith('### '):
            p = doc.add_heading('', level=2)
            add_formatted_text(p, line[4:].strip(), '黑体', 12, bold=True)
        elif line.startswith('#### '):
            p = doc.add_heading('', level=3)
            add_formatted_text(p, line[5:].strip(), '黑体', 11, bold=True)

        # 表格
        elif line.strip().startswith('|') and i + 1 < len(lines) and '---' in lines[i+1]:
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_text = lines[i].strip()
                if '---' not in row_text:
                    cells = [c.strip() for c in row_text.split('|')[1:-1]]
                    rows.append(cells)
                i += 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                for ri, row_cells in enumerate(rows):
                    for ci, cell_text in enumerate(row_cells):
                        cell = table.rows[ri].cells[ci]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        add_formatted_text(p, cell_text, '宋体', 9)
            continue

        # 引用块
        elif line.strip().startswith('> '):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip()[2:])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_text(p, '\n'.join(quote_lines), '楷体', 10)
            continue

        # 空行
        elif line.strip() == '':
            pass

        # 普通段落
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.3)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
            add_formatted_text(p, line.strip(), '宋体', 10.5)

        i += 1

    doc.save(docx_path)
    print(f"OK: {docx_path}")

def add_formatted_text(paragraph, text, font_name, size, bold=False):
    """添加格式化文本，支持**粗体**"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_chinese_font(run, font_name, size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_chinese_font(run, font_name, size, bold)

def add_title_page(doc, title, authors, abstract, keywords):
    """添加论文标题页"""
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_chinese_font(run, '黑体', 18, bold=True)
    doc.add_paragraph()

    # 作者
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(authors)
    set_chinese_font(run, '宋体', 10.5)
    doc.add_paragraph()

    # 摘要标签
    p = doc.add_paragraph()
    run = p.add_run('【摘要】')
    set_chinese_font(run, '黑体', 10.5, bold=True)
    run = p.add_run(abstract)
    set_chinese_font(run, '宋体', 10.5)
    p.paragraph_format.first_line_indent = Inches(0.3)

    # 关键词
    p = doc.add_paragraph()
    run = p.add_run('【关键词】')
    set_chinese_font(run, '黑体', 10.5, bold=True)
    run = p.add_run(keywords)
    set_chinese_font(run, '宋体', 10.5)
    p.paragraph_format.first_line_indent = Inches(0.3)

    doc.add_page_break()

if __name__ == '__main__':
    papers_dir = r'C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\papers'
    papers = [
        ('paper_15_anonymization.md', '论文15_匿名化标准_数据流通匿名化处理标准下的政府数据开放平台数据集质量评估.docx'),
        ('paper_03_data_sovereignty.md', '论文03_数据主权_总体国家安全观视域下政府数据开放平台的数据主权安全风险评估.docx'),
        ('paper_D_methodology.md', '论文D_方法论创新_政府数据开放绩效评估的方法论创新.docx'),
        ('paper_A_value_orientation.md', '论文A_价值导向_学科发展为了谁视域下政府数据开放的价值导向研究.docx'),
        ('paper_C_innovation.md', '论文C_守正创新_传统信息资源管理理论与数字时代数据治理的融合路径研究.docx'),
        ('paper_F_digital_china.md', '论文F_数字中国_政府数据开放赋能数字中国建设的路径研究.docx'),
        ('paper_B_concept.md', '论文B_概念谱系_政府数据开放的中国概念谱系.docx'),
        ('paper_E_ip_balance.md', '论文E_知识产权_公共数据开放中的开放悖论.docx'),
    ]

    for md_name, docx_name in papers:
        md_path = os.path.join(papers_dir, md_name)
        docx_path = os.path.join(papers_dir, docx_name)
        if os.path.exists(md_path):
            convert_markdown_to_docx(md_path, docx_path)
        else:
            print(f"MISSING: {md_path}")

    print("\nAll conversions completed!")
