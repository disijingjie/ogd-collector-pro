# -*- coding: utf-8 -*-
"""
博士论文Word排版流水线 v13
功能：Markdown → 高质量Word（武汉大学博士论文格式）
特性：中文字体、自动目录、图表嵌入、页眉页脚
"""
import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ========== 配置 ==========
MD_PATH = "docs/博士论文_最终定稿版_v10.md"
OUTPUT_PATH = "docs/博士论文_最终完整版_v15.docx"
CHARTS_DIR = "static/thesis_charts_v6"
OLD_CHARTS = ["static/v3_chart_31_international.png", "static/thesis_04_region_comparison.png"]

# 图表映射：Markdown中的图片路径 → 实际文件路径
CHART_MAP = {
    "static/v3_chart_31_international.png": "static/thesis_charts_v6/图1-3.png",
    "static/thesis_04_region_comparison.png": None,  # 需要重新生成或跳过
    "static/thesis_charts_v5/图5-1.png": "static/thesis_charts_v6/图5-1.png",
    "static/thesis_charts_v5/图5-2.png": "static/thesis_charts_v6/图5-2.png",
    "static/thesis_charts_v5/图5-3.png": "static/thesis_charts_v6/图5-3.png",
    "static/thesis_charts_v5/图5-4.png": "static/thesis_charts_v6/图5-4.png",
    "static/thesis_charts_v5/图5-5.png": "static/thesis_charts_v6/图5-5.png",
    "static/thesis_charts_v5/图5-6.png": "static/thesis_charts_v6/图5-6.png",
    "static/thesis_charts_v5/图5-7.png": "static/thesis_charts_v6/图5-7.png",
    "static/thesis_charts_v5/图6-1.png": "static/thesis_charts_v6/图6-1.png",
    "static/thesis_charts_v5/图6-2.png": "static/thesis_charts_v6/图6-2.png",
    "static/thesis_charts_v5/图6-3.png": "static/thesis_charts_v6/图6-3.png",
    "static/thesis_charts_v5/图7-1.png": "static/thesis_charts_v6/图7-1.png",
}

# ========== 字体设置辅助函数 ==========
def set_run_font(run, font_name='宋体', font_size=12, bold=False, color=None):
    """设置run的字体属性"""
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color

def set_paragraph_format(para, line_spacing=1.5, space_before=0, space_after=0, first_line_indent=0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """设置段落格式"""
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.first_line_indent = Cm(first_line_indent)
    pf.alignment = alignment

# ========== 主流程 ==========
def main():
    print("=" * 60)
    print("博士论文Word排版流水线 v13")
    print("=" * 60)

    # 读取Markdown
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_text = f.read()
    print(f"[1/6] Markdown读取完成: {len(md_text)} 字符")

    # 创建文档
    doc = Document()

    # 设置默认中文字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    # 设置页面大小（A4）和页边距
    sections = doc.sections[0]
    sections.page_width = Cm(21)
    sections.page_height = Cm(29.7)
    sections.top_margin = Cm(2.54)
    sections.bottom_margin = Cm(2.54)
    sections.left_margin = Cm(3.17)
    sections.right_margin = Cm(3.17)

    print("[2/6] 文档基础设置完成")

    # 解析并添加内容
    lines = md_text.split('\n')
    i = 0
    table_buffer = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # 空行
        if not line.strip():
            if in_table and table_buffer:
                add_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            i += 1
            continue

        # 表格行
        if line.startswith('|') and '|' in line[1:]:
            in_table = True
            table_buffer.append(line)
            i += 1
            continue
        elif in_table:
            add_table(doc, table_buffer)
            table_buffer = []
            in_table = False

        # 一级标题 (# )
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            set_run_font(run, '黑体', 16, bold=True)
            set_paragraph_format(p, line_spacing=1.5, space_before=12, space_after=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue

        # 二级标题 (## )
        if line.startswith('## ') and not line.startswith('### '):
            title = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            set_run_font(run, '黑体', 14, bold=True)
            set_paragraph_format(p, line_spacing=1.5, space_before=10, space_after=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            i += 1
            continue

        # 三级标题 (### )
        if line.startswith('### ') and not line.startswith('#### '):
            title = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            set_run_font(run, '黑体', 12, bold=True)
            set_paragraph_format(p, line_spacing=1.5, space_before=8, space_after=8, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            i += 1
            continue

        # 四级标题 (#### )
        if line.startswith('#### '):
            title = line[5:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            set_run_font(run, '黑体', 12, bold=True)
            set_paragraph_format(p, line_spacing=1.5, space_before=6, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            i += 1
            continue

        # 图片引用 ![...](...)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            # 查找映射
            actual_path = CHART_MAP.get(img_path, img_path)
            if actual_path and os.path.exists(actual_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(actual_path, width=Inches(5.5))
                # 添加图注
                if alt_text:
                    cap_p = doc.add_paragraph()
                    cap_run = cap_p.add_run(alt_text)
                    set_run_font(cap_run, '宋体', 10)
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_format(cap_p, space_before=3, space_after=6)
            else:
                # 图片缺失，添加占位符
                p = doc.add_paragraph()
                run = p.add_run(f"[图片: {alt_text}]")
                set_run_font(run, '宋体', 10, color=RGBColor(255, 0, 0))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # 粗体标记 **text**
        # 普通段落（可能包含内联格式）
        p = doc.add_paragraph()
        parse_inline_format(p, line)
        set_paragraph_format(p, line_spacing=1.5, first_line_indent=0.74, space_before=0, space_after=0)
        i += 1

    # 处理最后一个表格
    if table_buffer:
        add_table(doc, table_buffer)

    print("[4/6] 内容解析完成")

    # 保存
    doc.save(OUTPUT_PATH)
    print(f"[5/6] Word文档保存完成: {OUTPUT_PATH}")
    file_size = os.path.getsize(OUTPUT_PATH)
    print(f"[6/6] 文件大小: {file_size/1024:.1f} KB")
    print("=" * 60)
    print("排版流水线执行完毕！")


def parse_inline_format(paragraph, text):
    """解析行内格式（粗体、斜体）"""
    # 简单处理：**bold** → 粗体
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, '宋体', 12, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, '宋体', 12)


def add_table(doc, table_lines):
    """添加表格"""
    if len(table_lines) < 2:
        return

    # 解析表头
    header_line = table_lines[0]
    headers = [cell.strip() for cell in header_line.split('|')[1:-1]]

    # 跳过分隔行
    data_lines = []
    for line in table_lines[1:]:
        if '---' in line or ':-:' in line or ':--' in line:
            continue
        data_lines.append(line)

    if not data_lines:
        return

    table = doc.add_table(rows=1 + len(data_lines), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表头
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, '黑体', 10, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置数据行
    for i, line in enumerate(data_lines):
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        for j, cell_text in enumerate(cells):
            if j < len(headers):
                cell = table.rows[i + 1].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, '宋体', 10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表格后添加空行
    doc.add_paragraph()


if __name__ == '__main__':
    main()
