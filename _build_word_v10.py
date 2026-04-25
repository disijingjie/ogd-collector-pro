#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将v10 Markdown转换为Word文档"""
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx', '-q'])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_MD = Path('docs/博士论文_最终定稿版_v10.md')
OUTPUT_DOCX = Path('docs/博士论文_最终完整版_v10.docx')

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(1.18)
section.bottom_margin = Inches(0.98)
section.left_margin = Inches(1.10)
section.right_margin = Inches(0.98)

# 默认样式
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(28)
style.paragraph_format.first_line_indent = Inches(0.4)

content = INPUT_MD.read_text(encoding='utf-8')
lines = content.splitlines()

print(f"Total lines: {len(lines)}")

for line in lines:
    stripped = line.strip()
    if not stripped:
        continue
    
    # Chapter title
    if stripped.startswith('# ') and ('第' in stripped and '章' in stripped):
        p = doc.add_paragraph()
        run = p.add_run(stripped.replace('# ', ''))
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = '黑体'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
    
    # Section title (###)
    elif stripped.startswith('### '):
        p = doc.add_paragraph()
        run = p.add_run(stripped.replace('### ', ''))
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = '黑体'
        p.paragraph_format.first_line_indent = Inches(0)
    
    # Subsection (####)
    elif stripped.startswith('#### '):
        p = doc.add_paragraph()
        run = p.add_run(stripped.replace('#### ', ''))
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = '黑体'
        p.paragraph_format.first_line_indent = Inches(0)
    
    # Table header
    elif stripped.startswith('| ') and '---' not in stripped:
        p = doc.add_paragraph(stripped)
        p.paragraph_format.first_line_indent = Inches(0)
    
    # Image reference
    elif stripped.startswith('!['):
        # Extract caption
        m = re.match(r'!\[(.*?)\]', stripped)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(m.group(1))
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Inches(0)
    
    # Bold figure/table caption
    elif re.match(r'^\*\*(图|表)\d+-\d+', stripped):
        p = doc.add_paragraph()
        run = p.add_run(stripped.replace('**', ''))
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.name = '宋体'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
    
    # Regular text
    else:
        # Remove markdown bold markers for regular text
        text = stripped.replace('**', '')
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.line_spacing = Pt(28)

doc.save(str(OUTPUT_DOCX))
print(f"Word document saved: {OUTPUT_DOCX}")
print(f"File size: {OUTPUT_DOCX.stat().st_size / 1024:.1f} KB")
