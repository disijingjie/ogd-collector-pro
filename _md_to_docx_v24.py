# -*- coding: utf-8 -*-
"""将V24 Markdown转换为Word文档"""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx', '-q'])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_MD = Path('docs/新论文_完整版_V24_最终版.md')
OUTPUT_DOCX = Path('docs/博士论文_V24_最终版.docx')

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(1.18)
section.bottom_margin = Inches(0.98)
section.left_margin = Inches(1.10)
section.right_margin = Inches(0.98)

# 默认样式
default_style = doc.styles['Normal']
default_style.font.name = 'Times New Roman'
default_style.font.size = Pt(12)
default_style.paragraph_format.line_spacing = 1.5

try:
    default_style.element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '宋体')
except:
    pass

with open(INPUT_MD, 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.rstrip('\n')
    if not line.strip():
        continue

    if line.startswith('###### '):
        p = doc.add_heading(line[7:], level=6)
    elif line.startswith('##### '):
        p = doc.add_heading(line[6:], level=5)
    elif line.startswith('#### '):
        p = doc.add_heading(line[5:], level=4)
    elif line.startswith('### '):
        p = doc.add_heading(line[4:], level=3)
    elif line.startswith('## '):
        p = doc.add_heading(line[3:], level=2)
    elif line.startswith('# '):
        p = doc.add_heading(line[2:], level=1)
    elif line.startswith('> '):
        p = doc.add_paragraph(line[2:])
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
    elif line.startswith('|') and '|' in line[1:]:
        continue  # 跳过表格行
    elif line.startswith('---'):
        continue
    else:
        p = doc.add_paragraph(line)
        p.paragraph_format.first_line_indent = Inches(0.4)

doc.save(str(OUTPUT_DOCX))
print(f'Word文档已保存: {OUTPUT_DOCX}')

# 统计
import re
with open(INPUT_MD, 'r', encoding='utf-8') as f:
    content = f.read()
total_chars = len(re.sub(r'[\s#\-|=\n\r]', '', content))
chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', content))
print(f'总字符数: {total_chars:,}')
print(f'中文字符数: {chinese_chars:,}')
