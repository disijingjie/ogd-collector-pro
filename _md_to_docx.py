#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word converter for academic papers.
Uses python-docx. Handles headings, bold, tables, paragraphs.
"""
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_chinese_font(run, font_name='宋体', size=Pt(12), bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.font.bold = bold


def add_formatted_text(paragraph, text):
    """Parse inline formatting: **bold**, *italic*"""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_chinese_font(run, bold=True)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            set_chinese_font(run)
            run.font.italic = True
        else:
            run = paragraph.add_run(part)
            set_chinese_font(run)


def parse_table_row(line):
    """Parse a markdown table row into cells."""
    cells = [c.strip() for c in line.split('|')]
    # Remove empty first/last cells caused by leading/trailing |
    if cells and cells[0] == '':
        cells = cells[1:]
    if cells and cells[-1] == '':
        cells = cells[:-1]
    return cells


def is_table_separator(row):
    """Check if row is a markdown table separator like |---|---|"""
    cells = parse_table_row(row)
    for c in cells:
        if not re.match(r'^[:\s\-]+$', c):
            return False
    return True


def md_to_docx(md_path, docx_path):
    doc = Document()

    # Configure default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Empty line
        if not stripped:
            if in_table and table_rows:
                _flush_table(doc, table_rows)
                in_table = False
                table_rows = []
            i += 1
            continue

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            if in_table and table_rows:
                _flush_table(doc, table_rows)
                in_table = False
                table_rows = []
            p = doc.add_heading(level=1)
            run = p.add_run(stripped[2:])
            set_chinese_font(run, font_name='黑体', size=Pt(16), bold=True)
            i += 1
            continue

        if stripped.startswith('## '):
            if in_table and table_rows:
                _flush_table(doc, table_rows)
                in_table = False
                table_rows = []
            p = doc.add_heading(level=2)
            run = p.add_run(stripped[3:])
            set_chinese_font(run, font_name='黑体', size=Pt(14), bold=True)
            i += 1
            continue

        if stripped.startswith('### '):
            if in_table and table_rows:
                _flush_table(doc, table_rows)
                in_table = False
                table_rows = []
            p = doc.add_heading(level=3)
            run = p.add_run(stripped[4:])
            set_chinese_font(run, font_name='黑体', size=Pt(12), bold=True)
            i += 1
            continue

        # Table detection
        if '|' in stripped:
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(stripped)
            i += 1
            continue

        # Normal paragraph
        if in_table and table_rows:
            _flush_table(doc, table_rows)
            in_table = False
            table_rows = []

        # Horizontal rule
        if re.match(r'^\s*[-]{3,}\s*$', stripped):
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1

    if in_table and table_rows:
        _flush_table(doc, table_rows)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


def _flush_table(doc, rows):
    """Convert collected markdown table rows to a Word table."""
    # Filter out separator rows
    data_rows = [r for r in rows if not is_table_separator(r)]
    if not data_rows:
        return

    parsed = [parse_table_row(r) for r in data_rows]
    if not parsed:
        return

    num_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=num_cols)
    table.style = 'Table Grid'

    for r_idx, row_cells in enumerate(parsed):
        for c_idx in range(num_cols):
            cell = table.cell(r_idx, c_idx)
            if c_idx < len(row_cells):
                # Clear default paragraph, add formatted text
                p = cell.paragraphs[0]
                p.clear()
                add_formatted_text(p, row_cells[c_idx])


if __name__ == '__main__':
    files = [
        (r'C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文F_fsQCA组态分析_初稿.md',
         r'C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文F_fsQCA组态分析_投稿版.docx'),
        (r'C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文B_数据口径幻觉_初稿.md',
         r'C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文B_数据口径幻觉_投稿版.docx'),
        (r'C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文D_TOPSIS-DEA绩效评估_初稿.md',
         r'C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文D_TOPSIS-DEA绩效评估_投稿版.docx'),
        (r'C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文A_异构平台统一采集_初稿.md',
         r'C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文A_异构平台统一采集_投稿版.docx'),
        (r'C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文C_4E评估框架_初稿.md',
         r'C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\论文C_4E评估框架_投稿版.docx'),
    ]

    for md_path, docx_path in files:
        try:
            md_to_docx(md_path, docx_path)
        except Exception as e:
            print(f"Error converting {md_path}: {e}")
