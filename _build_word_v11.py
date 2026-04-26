#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将v11 Markdown转换为Word文档 - 改进版
改进点：
1. 正确解析Markdown表格为Word表格
2. 图片引用尝试插入实际图片
3. 更好的标题层级处理
4. 支持脚注和注释
"""
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx', '-q'])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

INPUT_MD = Path('docs/博士论文_最终定稿版_v10.md')
OUTPUT_DOCX = Path('docs/博士论文_最终完整版_v11.docx')
CHARTS_DIR = Path('static/thesis_charts_v4')

doc = Document()

# 页面设置 - A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# 默认样式
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(28)
style.paragraph_format.first_line_indent = Cm(0.74)

content = INPUT_MD.read_text(encoding='utf-8')
lines = content.splitlines()

print(f"Total lines: {len(lines)}")

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    
    if not stripped:
        i += 1
        continue
    
    # Chapter title (# 第一章)
    if stripped.startswith('# ') and ('章' in stripped):
        p = doc.add_paragraph()
        run = p.add_run(stripped.replace('# ', ''))
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(18)
    
    # Section title (##)
    elif stripped.startswith('## '):
        p = doc.add_paragraph()
        run = p.add_run(stripped.replace('## ', ''))
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
    
    # Subsection (###)
    elif stripped.startswith('### '):
        p = doc.add_paragraph()
        run = p.add_run(stripped.replace('### ', ''))
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    
    # Subsubsection (####)
    elif stripped.startswith('#### '):
        p = doc.add_paragraph()
        run = p.add_run(stripped.replace('#### ', ''))
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
    
    # Table
    elif stripped.startswith('| ') and '---' not in stripped:
        # Collect table rows
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        
        # Skip separator line
        if len(table_lines) >= 2 and '---' in table_lines[1]:
            header_line = table_lines[0]
            data_lines = table_lines[2:]
        else:
            header_line = table_lines[0]
            data_lines = table_lines[1:]
        
        # Parse cells
        def parse_row(line):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            return cells
        
        headers = parse_row(header_line)
        if not headers:
            continue
        
        # Create table
        table = doc.add_table(rows=1+len(data_lines), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        for j, h in enumerate(headers):
            hdr_cells[j].text = h.replace('**', '')
            for paragraph in hdr_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10.5)
                    run.font.name = '宋体'
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for row_idx, data_line in enumerate(data_lines):
            cells = parse_row(data_line)
            row_cells = table.rows[row_idx+1].cells
            for j, c in enumerate(cells):
                if j < len(row_cells):
                    row_cells[j].text = c.replace('**', '')
                    for paragraph in row_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10.5)
                            run.font.name = '宋体'
                        # Right-align numbers
                        if c.replace('.', '').replace('-', '').isdigit():
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add spacing after table
        doc.add_paragraph()
        continue  # i already advanced
    
    # Image reference - try to insert actual image
    elif stripped.startswith('!['):
        m = re.match(r'!\[(.*?)\]\((.*?)\)', stripped)
        if m:
            caption, img_path = m.groups()
            # Try to find image
            img_file = CHARTS_DIR / Path(img_path).name
            if not img_file.exists():
                # Try alternative paths
                alt_paths = [
                    Path(img_path),
                    CHARTS_DIR / img_path.replace('static/', ''),
                    CHARTS_DIR / img_path.replace('static/thesis_charts/', ''),
                ]
                for alt in alt_paths:
                    if alt.exists():
                        img_file = alt
                        break
            
            if img_file.exists():
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_file), width=Cm(14))
                except Exception as e:
                    print(f"  Warning: Could not insert image {img_file}: {e}")
                    p = doc.add_paragraph()
                    run = p.add_run(f"[图片: {caption}]")
                    run.font.size = Pt(10.5)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[图片: {caption}]")
                run.font.size = Pt(10.5)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Bold figure/table caption
    elif re.match(r'^\*\*(图|表)\d+-\d+', stripped):
        p = doc.add_paragraph()
        run = p.add_run(stripped.replace('**', ''))
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.name = '宋体'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(6)
    
    # Italic note
    elif stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(stripped.replace('*', ''))
        run.font.size = Pt(10.5)
        run.font.italic = True
        run.font.name = '宋体'
        p.paragraph_format.first_line_indent = Cm(0)
    
    # Regular text with bold markers
    else:
        # Handle inline bold
        text = stripped
        if text:
            p = doc.add_paragraph()
            # Split by bold markers
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(12)
                run.font.name = '宋体'
            p.paragraph_format.line_spacing = Pt(28)
    
    i += 1

doc.save(str(OUTPUT_DOCX))
print(f"Word document saved: {OUTPUT_DOCX}")
print(f"File size: {OUTPUT_DOCX.stat().st_size / 1024:.1f} KB")
