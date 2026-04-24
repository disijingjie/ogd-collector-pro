# -*- coding: utf-8 -*-
"""
逐章提取论文内容，保存为独立文件供深度分析
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
import re

doc = Document('D:/Users/MI/120599705/WPS云盘/博士毕业/【排版结果】博士论文_V14_精修版.docx')

# 章节标记模式
chapter_markers = [
    ("摘要", "abstract"),
    ("Abstract", "abstract_en"),
    ("绪论", "introduction"),
    ("第一章", "ch1"),
    ("第二章", "ch2"),
    ("第三章", "ch3"),
    ("第四章", "ch4"),
    ("第五章", "ch5"),
    ("第六章", "ch6"),
    ("第七章", "ch7"),
    ("第八章", "ch8"),
    ("结论与展望", "conclusion"),
    ("参考文献", "references"),
    ("致谢", "acknowledgement"),
    ("附录", "appendix"),
]

# 找到各章起始位置
chapter_starts = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    for marker, code in chapter_markers:
        if text.startswith(marker) or (marker in text[:20] and len(text) < 100):
            chapter_starts.append((i, code, text[:80]))
            break

print(f"找到 {len(chapter_starts)} 个章节标记点")
for idx, (pos, code, title) in enumerate(chapter_starts):
    print(f"  [{idx}] {code}: 段落{pos} - {title}")

# 提取各章内容
for idx in range(len(chapter_starts)):
    start_pos = chapter_starts[idx][0]
    code = chapter_starts[idx][1]
    if idx + 1 < len(chapter_starts):
        end_pos = chapter_starts[idx + 1][0]
    else:
        end_pos = len(doc.paragraphs)
    
    # 提取该章所有段落
    chapter_texts = []
    for i in range(start_pos, end_pos):
        text = doc.paragraphs[i].text.strip()
        if text:
            chapter_texts.append(text)
    
    # 保存到文件
    filename = f"_thesis_{code}.txt"
    with open(filename, 'w', encoding='utf-8') as f:
        f.write('\n'.join(chapter_texts))
    
    total_chars = sum(len(t) for t in chapter_texts)
    print(f"  已保存 {code}: {len(chapter_texts)} 段落, {total_chars} 字符 -> {filename}")

print("\n全部章节提取完成")
