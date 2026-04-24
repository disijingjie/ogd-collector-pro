# -*- coding: utf-8 -*-
"""
逐章提取论文内容 v2 - 通过Word样式识别章节结构
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
import re

doc = Document('D:/Users/MI/120599705/WPS云盘/博士毕业/【排版结果】博士论文_V14_精修版.docx')

print(f"总段落数: {len(doc.paragraphs)}")
print(f"总表格数: {len(doc.tables)}")

# 分析前500个段落的样式，找出章节标题的样式模式
print("\n=== 分析段落样式分布（前500段）===")
style_counts = {}
for i in range(min(500, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    style_name = para.style.name if para.style else "None"
    text = para.text.strip()
    if style_name not in style_counts:
        style_counts[style_name] = 0
    style_counts[style_name] += 1

for style, count in sorted(style_counts.items(), key=lambda x: -x[1]):
    print(f"  {style}: {count} 段落")

# 分析全部段落的样式
print("\n=== 全部段落样式分布 ===")
all_style_counts = {}
heading_styles = []
for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else "None"
    if style_name not in all_style_counts:
        all_style_counts[style_name] = 0
    all_style_counts[style_name] += 1
    
    # 检测可能的标题样式
    text = para.text.strip()
    if text and len(text) < 200:
        # 标题样式关键词
        if any(k in style_name.lower() for k in ['heading', '标题', 'heading', 'title', '题']):
            heading_styles.append((i, style_name, text[:100]))
        # 或者文本内容看起来像标题
        elif re.match(r'^(第[一二三四五六七八九十]+章|第\d+章|绪论|结论|参考文献|致谢|附录)', text):
            heading_styles.append((i, style_name, text[:100]))

for style, count in sorted(all_style_counts.items(), key=lambda x: -x[1])[:20]:
    print(f"  {style}: {count} 段落")

print(f"\n=== 检测到 {len(heading_styles)} 个可能的标题段落 ===")
for i, (pos, style, text) in enumerate(heading_styles[:50]):
    print(f"  [{i}] 段落{pos} | 样式:{style} | {text}")
