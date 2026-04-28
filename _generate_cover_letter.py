#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate Cover Letter templates for 5 papers."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


def set_font(run, name='Times New Roman', east_asia='宋体', size=Pt(12), bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    run.font.size = size
    run.font.bold = bold


def add_paragraph(doc, text, bold=False, italic=False, align=None, font_size=Pt(12), east_asia='宋体'):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=font_size, bold=bold, east_asia=east_asia)
    if italic:
        run.font.italic = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    return p


def generate_cover_letter(paper_title, paper_type, journal_name, highlights, fit_reason, output_path):
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Date
    add_paragraph(doc, "2026年4月28日", align=WD_ALIGN_PARAGRAPH.RIGHT, font_size=Pt(10.5))

    # Salutation
    add_paragraph(doc, f"尊敬的《{journal_name}》编辑部：", bold=True, font_size=Pt(12), east_asia='黑体')

    # Opening
    add_paragraph(doc,
        f"我们谨向贵刊投稿一篇{paper_type}论文，论文题目为《{paper_title}》。"
        f"该论文为未公开发表的原创性研究成果，未以任何语言在任何刊物上发表，也未一稿多投。")

    # Innovation highlights
    add_paragraph(doc, "本研究的主要创新点包括：", bold=True, font_size=Pt(12), east_asia='黑体')
    for idx, hl in enumerate(highlights, 1):
        add_paragraph(doc, f"{idx}. {hl}")

    # Fit reason
    add_paragraph(doc, "投稿适配性说明：", bold=True, font_size=Pt(12), east_asia='黑体')
    add_paragraph(doc, fit_reason)

    # Suggested reviewers (optional)
    add_paragraph(doc, "建议审稿人（如有需要）：", bold=True, font_size=Pt(12), east_asia='黑体')
    add_paragraph(doc, "由于我们对本研究领域国内外学者的了解有限，恳请编辑部根据论文主题指派合适的审稿专家。")

    # Conflict of interest
    add_paragraph(doc, "利益冲突声明：", bold=True, font_size=Pt(12), east_asia='黑体')
    add_paragraph(doc, "本研究不存在任何利益冲突。研究数据来源于公开的政府数据开放平台，不涉及人类受试者或伦理审查问题。")

    # Closing
    add_paragraph(doc, "恳请贵刊予以审稿。如需任何补充材料或信息，请随时与我们联系。")
    add_paragraph(doc, "此致", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "敬礼！", align=WD_ALIGN_PARAGRAPH.LEFT)

    # Signature block
    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "投稿人：文明", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "武汉大学信息管理学院", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "电子邮箱：ambit@qq.com", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "通讯地址：湖北省武汉市武昌区珞珈山武汉大学信息管理学院", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "邮政编码：430072", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "日期：2026年4月28日", align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(output_path)
    print(f"Generated: {output_path}")


PAPERS = [
    {
        "title": "异构政府数据开放平台的统一采集与标准化方法研究——基于23个省级平台的'一平台一策'实践",
        "type": "研究方法",
        "journal": "数据分析与知识发现",
        "highlights": [
            "自主开发了OGD-Collector Pro三层架构采集系统，实现了对31个省级平台的差异化采集；",
            "提出了'概念等价性'替代'技术同质性'的方法论思路，构建了'采集-标准化-验证'的三层方法论框架；",
            "建立了可操作的口径转换系数体系和置信度分级标准，可直接应用于其他异构数据采集场景；",
            "首次对全国23个省级平台的'数据口径幻觉'进行了系统性测量，揭示了口径一致性系数中位数仅为0.122的严峻现实。"
        ],
        "fit": "《数据分析与知识发现》聚焦'以大数据为基础、依靠数据挖掘分析、进行知识发现与预测'的研究，广泛吸纳数据科学、情报科学与计算机科学领域的技术与方法。本研究涉及的异构平台自动化采集、数据标准化、口径转换与置信度评估，与该刊的核心方向高度契合。",
        "output": r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\CoverLetter_论文A_数据分析与知识发现.docx"
    },
    {
        "title": "'数据口径幻觉'：政府数据开放评估中的测量偏差与行为扭曲",
        "type": "理论研究",
        "journal": "情报学报",
        "highlights": [
            "首次提出并操作化'数据口径幻觉'概念，填补了政府数据开放评估中测量偏差研究的理论空白；",
            "将Goodhart法则和Campbell定律引入政府数据开放评估领域，构建了'指标设计-行为激励-绩效失真'的理论分析链条；",
            "基于23个省级平台的实地采集数据，发现口径一致性系数中位数仅为0.122，山东平台低至0.016；",
            "揭示了评估体系作为一种'选择架构'对平台运营者行为的系统性引导作用，为公共部门绩效评估理论提供了新的经验证据。"
        ],
        "fit": "《情报学报》作为中国科学技术情报学会会刊，长期关注情报科学领域的方法创新与理论前沿。本研究提出的'数据口径幻觉'概念及其测量方法，是对政府数据开放评估理论的原创性贡献，符合该刊对'情报学理论创新'和'科学测量方法'的选稿偏好。",
        "output": r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\CoverLetter_论文B_情报学报.docx"
    },
    {
        "title": "效果导向的政府数据开放评估框架构建——基于4E框架的迁移与适配",
        "type": "理论框架",
        "journal": "信息资源管理学报",
        "highlights": [
            "基于4E绩效评估框架，结合公共价值理论和数据要素价值化理论，构建了涵盖5个一级维度、24个具体指标的评估体系；",
            "将'利用效果'维度权重提升至30%，引入了授权运营成效、API调用深度等前瞻性指标，推动评估范式从'供给导向'向'效果导向'转型；",
            "通过与开放数林指数、ODI指数等主流评估工具的系统比较，揭示了现有评估体系'重供给轻效果'的结构性局限；",
            "建立了从'数据资源化'到'数据资本化'的全链条价值映射，为政府数据开放评估提供了新的理论视角。"
        ],
        "fit": "《信息资源管理学报》由武汉大学主办，是信息资源管理领域与学科名称相同的权威学术期刊。本研究作者文明系武汉大学信息管理学院博士研究生，论文主题'政府数据开放评估框架'与期刊的'信息开发与利用'、'信息经济与政策'等核心栏目高度契合。投稿贵刊既符合学科归属，也有助于深化母校学科的理论建设。",
        "output": r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\CoverLetter_论文C_信息资源管理学报.docx"
    },
    {
        "title": "中国省级政府数据开放平台绩效评估——基于TOPSIS-DEA的实证研究",
        "type": "实证研究",
        "journal": "管理评论",
        "highlights": [
            "构建了'绩效-效率'二维分析框架，发现23个平台中无一个达到DEA有效，揭示了'高绩效低效率'悖论；",
            "运用TOPSIS-DEA组合方法，既评估绩效水平（'是什么'），又评估效率高低（'是否高效'），弥补了现有评估工具效率维度的缺失；",
            "通过绩效-效率矩阵将平台分为四类，为差异化政策干预提供了精准靶向；",
            "发现区域差异显著，但东部内部差异同样突出，挑战了简单的'东中西部'二元叙事。"
        ],
        "fit": "《管理评论》是中国科学院大学主办的经济管理科学权威期刊，注重管理科学理论与方法的创新应用。本研究采用的TOPSIS-DEA组合评估方法属于管理科学经典方法论的交叉融合，绩效-效率二维分析框架对公共部门资源配置研究具有方法论示范价值，符合该刊对'管理科学方法创新'的选稿导向。",
        "output": r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\CoverLetter_论文D_管理评论.docx"
    },
    {
        "title": "高绩效政府数据开放平台的多重等效路径——基于fsQCA的组态分析",
        "type": "实证研究",
        "journal": "公共管理学报",
        "highlights": [
            "将fsQCA方法系统引入政府数据开放研究领域，推动该领域从'线性思维'向'组态思维'的方法论转型；",
            "发现了三条高绩效组态路径（全要素驱动型、服务-质量-效果驱动型、生态协同型），验证了fsQCA的等价性原则；",
            "揭示了低绩效平台的'供给-质量双低型'和'制度-技术双低型'两种组态，印证了fsQCA的非对称性原则；",
            "通过DEMATEL-fsQCA交叉验证，增强了研究结论的稳健性，为公共管理领域的'多因一果'现象提供了方法论示范。"
        ],
        "fit": "《公共管理学报》由哈尔滨工业大学主办，是公共管理领域的重要学术期刊，坚持'求道无篱，经世致用'的办刊原则，注重以事实为基础的实证研究和中国场景下的问题导向。本研究将fsQCA组态分析方法应用于政府数据开放绩效分析，既回应了'不同资源禀赋平台如何高绩效'的现实问题，也为公共管理研究提供了方法论创新，与该刊的办刊理念高度一致。",
        "output": r"C:\Users\MI\WorkBuddy\newbbbb\ogd_collector_system\docs\CoverLetter_论文F_公共管理学报.docx"
    },
]

if __name__ == '__main__':
    for p in PAPERS:
        generate_cover_letter(
            paper_title=p["title"],
            paper_type=p["type"],
            journal_name=p["journal"],
            highlights=p["highlights"],
            fit_reason=p["fit"],
            output_path=p["output"]
        )
